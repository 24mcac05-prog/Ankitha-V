# app_with_pinecone_ollama.py
"""
Jarvis — Streamlit document assistant (local with optional Pinecone + Ollama/TinyLlama)
- Single-file upload -> local TF-IDF retriever (fallback)
- Optional Pinecone: upsert/search using sentence-transformers embeddings
- Optional Ollama: use local LLM (tinyllama/phi3) for nicer summarization (but answers must still come from uploaded docs)
- Simple SQLite auth (email/password, PBKDF2) for local/dev only
- Streamlit UI: upload single file, ask questions (answers only from uploaded doc),
  summarizer (5 lines), chat history, export to PDF/DOCX, New Chat.
"""

import os
from pathlib import Path
import uuid
import io
import threading
import tempfile
import time
import re
import unicodedata
import math
import sqlite3
import hashlib
import binascii
import json
import sys
import subprocess

import streamlit as st

# --- Optional imports (safe tries) ---
# PDF / DOCX parsing
try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx
except Exception:
    docx = None

# TTS / audio
USE_SOUNDDEVICE = False
USE_PYTTSX3 = False
try:
    import sounddevice as sd
    import wavio
    USE_SOUNDDEVICE = True
except Exception:
    USE_SOUNDDEVICE = False

try:
    import pyttsx3
    USE_PYTTSX3 = True
except Exception:
    USE_PYTTSX3 = False

# sentence-transformers (for Pinecone embeddings)
try:
    from sentence_transformers import SentenceTransformer
except Exception:
    SentenceTransformer = None

# Pinecone modern client
try:
    from pinecone import Pinecone, ServerlessSpec
except Exception:
    Pinecone = None
    ServerlessSpec = None

# Ollama local LLM wrapper (if available)
try:
    from langchain_community.llms import Ollama
    OllamaAvailable = True
except Exception:
    OllamaAvailable = False
    Ollama = None

# FPDF for PDF export
try:
    from fpdf import FPDF
except Exception:
    FPDF = None

# python-docx import name for writing
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

# ------------------- Streamlit config -------------------
st.set_page_config(page_title="Jarvis — Pinecone + TinyLlama optional", layout="wide", initial_sidebar_state="expanded")

# ------------------- helpers -------------------
def safe_rerun():
    if hasattr(st, "rerun"):
        try:
            st.rerun()
            return
        except Exception:
            pass
    if hasattr(st, "experimental_rerun"):
        try:
            st.experimental_rerun()
            return
        except Exception:
            pass
    st.session_state["_force_update_flag"] = not st.session_state.get("_force_update_flag", False)

def human_readable_size(num_bytes: int) -> str:
    for unit in ["B", "KB", "MB", "GB", "TB"]:
        if abs(num_bytes) < 1024.0:
            return f"{num_bytes:3.1f} {unit}"
        num_bytes /= 1024.0
    return f"{num_bytes:.1f} PB"

def safe_for_pdf(s: str) -> str:
    if s is None:
        return ""
    if not isinstance(s, str):
        s = str(s)
    normalized = unicodedata.normalize("NFKD", s)
    return normalized.encode("latin-1", "ignore").decode("latin-1")

def make_concise_paragraph(snippets, max_sentences=3, max_chars=420):
    if not snippets:
        return ""
    merged = []
    for s in snippets:
        if not s:
            continue
        t = s.strip().replace("\n", " ")
        t = re.sub(r"[“”«»]", '"', t)
        t = re.sub(r"[‘’]", "'", t)
        if len(t) < 3:
            continue
        merged.append(t)
    if not merged:
        return ""
    first = merged[0]
    extend_needed = False
    if first.endswith(":") or first.endswith("?") or len(first.split()) <= 6:
        extend_needed = True
    if extend_needed and len(merged) > 1:
        combined = " ".join(merged[: min(len(merged), 4)])
    else:
        combined = " ".join(merged)
    sentences = re.split(r'(?<=[.!?])\s+', combined)
    sentences = [s.strip() for s in sentences if s.strip()]
    if not sentences:
        paragraph = combined.strip()
    else:
        paragraph = " ".join(sentences[:max_sentences])
    paragraph = re.sub(r"•\s*", "", paragraph)
    paragraph = re.sub(r"\s+", " ", paragraph).strip()
    if len(paragraph) > max_chars:
        paragraph = paragraph[: max_chars - 3].rstrip() + "..."
    return paragraph

# ------------------- Auth (sqlite + PBKDF2) -------------------
USERS_DB_PATH = Path(__file__).resolve().parent / "users.db"

def init_user_db():
    conn = sqlite3.connect(str(USERS_DB_PATH))
    cur = conn.cursor()
    cur.execute("""
      CREATE TABLE IF NOT EXISTS users (
        email TEXT PRIMARY KEY,
        salt TEXT NOT NULL,
        pw_hash TEXT NOT NULL,
        created_at INTEGER NOT NULL
      )
    """)
    conn.commit()
    conn.close()

def hash_password(password: str, salt: bytes = None):
    if salt is None:
        salt = os.urandom(16)
    dk = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, 100_000)
    return binascii.hexlify(salt).decode("ascii"), binascii.hexlify(dk).decode("ascii")

def create_user(email: str, password: str) -> (bool, str):
    email = (email or "").strip().lower()
    if not email or not password:
        return False, "Email and password required."
    conn = sqlite3.connect(str(USERS_DB_PATH))
    cur = conn.cursor()
    cur.execute("SELECT 1 FROM users WHERE email = ?", (email,))
    if cur.fetchone():
        conn.close()
        return False, "Email already registered."
    salt_hex, hash_hex = hash_password(password)
    ts = int(time.time())
    try:
        cur.execute("INSERT INTO users (email, salt, pw_hash, created_at) VALUES (?, ?, ?, ?)",
                    (email, salt_hex, hash_hex, ts))
        conn.commit()
    except Exception as e:
        conn.close()
        return False, f"DB error: {e}"
    conn.close()
    return True, "Account created."

def verify_user(email: str, password: str) -> (bool, str):
    email = (email or "").strip().lower()
    conn = sqlite3.connect(str(USERS_DB_PATH))
    cur = conn.cursor()
    cur.execute("SELECT salt, pw_hash FROM users WHERE email = ?", (email,))
    row = cur.fetchone()
    conn.close()
    if not row:
        return False, "No account with that email."
    salt_hex, stored_hash_hex = row
    salt = binascii.unhexlify(salt_hex.encode("ascii"))
    _, hash_hex = hash_password(password, salt=salt)
    if hash_hex == stored_hash_hex:
        return True, "Login successful."
    return False, "Incorrect password."

init_user_db()

# ------------------- Document parsers / indexing -------------------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    if not PyPDF2:
        return ""
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    pages = []
    for p in reader.pages:
        try:
            raw = p.extract_text() or ""
        except Exception:
            raw = ""
        lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
        merged = []
        buffer = ""
        for line in lines:
            if re.search(r'[.!?]$', line):
                buffer += " " + line
                merged.append(buffer.strip())
                buffer = ""
            else:
                buffer += " " + line
        if buffer.strip():
            merged.append(buffer.strip())
        pages.append(" ".join(merged))
    text = "\n".join(pages)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'•', '-', text)
    return text.strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    if not docx:
        return ""
    tmp = Path(tempfile.gettempdir()) / f"tmp_{uuid.uuid4().hex}.docx"
    tmp.write_bytes(file_bytes)
    document = docx.Document(str(tmp))
    paragraphs = [p.text for p in document.paragraphs]
    tmp.unlink(missing_ok=True)
    return "\n".join(paragraphs)

def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 150):
    chunks = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + chunk_size, length)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start += chunk_size - overlap
    return chunks

def index_chunks_in_session(chunks, source_name):
    for c in chunks:
        st.session_state.documents.append({"source": source_name, "text": c})

# ------------------- TF-IDF retriever (fallback) -------------------
def build_tfidf_index():
    docs = st.session_state.documents
    N = len(docs)
    if N == 0:
        return [], {}
    doc_tokens = []
    df = {}
    for d in docs:
        tokens = re.findall(r"\w+", d["text"].lower())
        uniq = set(tokens)
        for t in uniq:
            df[t] = df.get(t, 0) + 1
        doc_tokens.append(tokens)
    idf = {}
    for t, cnt in df.items():
        idf[t] = math.log((N + 1) / (cnt + 1)) + 1.0
    docs_vectors = []
    for tokens, d in zip(doc_tokens, docs):
        tf = {}
        for t in tokens:
            tf[t] = tf.get(t, 0) + 1
        vec = {}
        norm_sq = 0.0
        for t, tfcount in tf.items():
            w = tfcount * idf.get(t, 0.0)
            vec[t] = w
            norm_sq += w * w
        norm = math.sqrt(norm_sq) if norm_sq > 0 else 1.0
        docs_vectors.append({"doc": d, "vec": vec, "norm": norm})
    return docs_vectors, idf

def retrieve_from_session_tfidf(query: str, k: int = 3):
    if not st.session_state.documents or not query or not query.strip():
        return []
    q_tokens = re.findall(r"\w+", query.lower())
    if not q_tokens:
        return []
    docs_vectors, idf = build_tfidf_index()
    if not docs_vectors:
        return []
    q_tf = {}
    for t in q_tokens:
        q_tf[t] = q_tf.get(t, 0) + 1
    q_vec = {}
    q_norm_sq = 0.0
    for t, tfcount in q_tf.items():
        idf_val = idf.get(t, math.log((len(docs_vectors) + 1) / 1) + 1.0)
        w = tfcount * idf_val
        q_vec[t] = w
        q_norm_sq += w * w
    q_norm = math.sqrt(q_norm_sq) if q_norm_sq > 0 else 1.0
    scored = []
    q_text_lower = query.lower()
    for item in docs_vectors:
        doc = item["doc"]
        doc_vec = item["vec"]
        dot = 0.0
        for t, qw in q_vec.items():
            dw = doc_vec.get(t)
            if dw:
                dot += qw * dw
        denom = (q_norm * item["norm"]) if (q_norm and item["norm"]) else 1.0
        score = dot / denom if denom != 0 else 0.0
        if q_text_lower in doc["text"].lower():
            score += 0.18
        common_tokens = sum(1 for t in set(q_tokens) if t in doc["text"].lower())
        score += min(common_tokens * 0.02, 0.06)
        scored.append((score, doc))
    scored.sort(key=lambda x: x[0], reverse=True)
    top = []
    for s, d in scored:
        if s <= 0:
            continue
        top.append(d)
        if len(top) >= k:
            break
    return top

# ------------------- Pinecone integration (optional) -------------------
PINECONE_ENABLED = False
pinecone_client = None
pine_index = None
embed_model = None

def init_pinecone_if_configured():
    global PINECONE_ENABLED, pinecone_client, pine_index, embed_model
    api_key = os.getenv("PINECONE_API_KEY") or ""
    index_name = os.getenv("PINECONE_INDEX", "jarvis-index")
    # require sentence transformer for embeddings
    if not api_key:
        return
    if Pinecone is None:
        return
    try:
        # create client
        pinecone_client = Pinecone(api_key=api_key)
        # create index if missing (spec optional)
        try:
            # check list indexes API - new client returns an object; call .list_indexes()
            existing = pinecone_client.list_indexes()
            # Some distributions return a list directly, some return an object with .names(); handle both
            names = existing.names() if hasattr(existing, "names") else list(existing)
            if index_name not in names:
                # if ServerlessSpec available, create sample spec with default cloud/region
                kwargs = {"name": index_name, "dimension": 1536, "metric": "cosine"}
                if ServerlessSpec is not None:
                    # choose defaults; user can adjust environment variables to match their cloud
                    try:
                        spec = ServerlessSpec(cloud="aws", region="us-east-1")
                        pinecone_client.create_index(name=index_name, dimension=1536, metric="cosine", spec=spec)
                    except Exception:
                        pinecone_client.create_index(name=index_name, dimension=1536, metric="cosine")
                else:
                    pinecone_client.create_index(name=index_name, dimension=1536, metric="cosine")
            pine_index = pinecone_client.Index(index_name)
            # initialize embedder if available
            if SentenceTransformer is not None:
                try:
                    embed_model = SentenceTransformer("all-MiniLM-L6-v2")
                except Exception:
                    embed_model = None
            PINECONE_ENABLED = True
            st.session_state["pinecone_ok"] = True
            print("✅ Pinecone initialized (client created).")
        except Exception as e:
            print("⚠ Pinecone client created but index/create failed:", e)
            st.session_state["pinecone_ok"] = False
    except Exception as e:
        print("⚠ Pinecone init failed:", e)
        st.session_state["pinecone_ok"] = False

# try init at start (non-blocking)
init_pinecone_if_configured()

# Upsert helper (only if Pinecone enabled & embed_model available)
def upsert_to_pinecone(chunks, source_name):
    global pine_index, embed_model
    if not PINECONE_ENABLED or pine_index is None or embed_model is None:
        return False
    try:
        # embed in batches
        batch = []
        for i, c in enumerate(chunks):
            vec = embed_model.encode(c).tolist()
            vid = f"{source_name}_{uuid.uuid4().hex}_{i}"
            meta = {"source": source_name, "text": c}
            batch.append((vid, vec, meta))
            if len(batch) >= 50:
                pine_index.upsert(vectors=batch)
                batch = []
        if batch:
            pine_index.upsert(vectors=batch)
        return True
    except Exception as e:
        print("Pinecone upsert error:", e)
        return False

def query_pinecone(query, top_k=3):
    global pine_index, embed_model
    if not PINECONE_ENABLED or pine_index is None or embed_model is None:
        return []
    try:
        qvec = embed_model.encode(query).tolist()
        resp = pine_index.query(qvec, top_k=top_k, include_metadata=True)
        # Different client versions respond differently - try to normalize
        results = []
        if hasattr(resp, "matches"):
            matches = resp.matches
        elif isinstance(resp, dict) and "matches" in resp:
            matches = resp["matches"]
        else:
            matches = []
        for m in matches:
            # metadata may be in m.metadata or m["metadata"]
            meta = getattr(m, "metadata", None) or m.get("metadata", {}) if isinstance(m, dict) else {}
            txt = meta.get("text") or meta.get("content") or ""
            src = meta.get("source", "unknown")
            results.append({"source": src, "text": txt})
        return results
    except Exception as e:
        print("Pinecone query error:", e)
        return []

# ------------------- Ollama integration (optional) ----
LLM_ENABLED = False
llm_client = None
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "tinyllama")
def init_ollama_if_available():
    global LLM_ENABLED, llm_client, OllamaAvailable, Ollama, OLLAMA_MODEL
    if not OllamaAvailable:
        LLM_ENABLED = False
        st.session_state["llm_ok"] = False
        return
    try:
        llm_client = Ollama(model=OLLAMA_MODEL)
        LLM_ENABLED = True
        st.session_state["llm_ok"] = True
        print("✅ Ollama LLM client ready:", OLLAMA_MODEL)
    except Exception as e:
        print("⚠ Ollama load failed:", e)
        LLM_ENABLED = False
        st.session_state["llm_ok"] = False

init_ollama_if_available()

# Use llm_client.generate / .run only for optional summarization; the retrieval answers always come from uploaded docs

# ------------------- Session state init -------------------
if "user_email" not in st.session_state:
    st.session_state["user_email"] = None
if "auth_message" not in st.session_state:
    st.session_state["auth_message"] = ""
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "documents" not in st.session_state:
    st.session_state.documents = []
if "last_processed_query" not in st.session_state:
    st.session_state.last_processed_query = None
if "_cur_submission_token" not in st.session_state:
    st.session_state["_cur_submission_token"] = None
if "_handled_submission_token" not in st.session_state:
    st.session_state["_handled_submission_token"] = None

# ------------------- UI: Sidebar auth -------------------
with st.sidebar:
    st.title("🔐 Account")
    if st.session_state.get("user_email"):
        st.markdown(f"**Signed in as:** {st.session_state['user_email']}")
        if st.button("🔓 Logout"):
            st.session_state["user_email"] = None
            st.session_state["auth_message"] = "Logged out."
            try:
                safe_rerun()
            except Exception:
                pass
    else:
        auth_mode = st.radio("Choose:", ["Login", "Sign up"], index=0)
        email_in = st.text_input("Email", key="auth_email")
        pw_in = st.text_input("Password", type="password", key="auth_pw")
        if auth_mode == "Sign up":
            if st.button("Create account"):
                ok, msg = create_user(email_in, pw_in)
                st.session_state["auth_message"] = msg
                if ok:
                    st.success(msg)
                else:
                    st.error(msg)
        else:  # Login
            if st.button("Login"):
                ok, msg = verify_user(email_in, pw_in)
                st.session_state["auth_message"] = msg
                if ok:
                    st.session_state["user_email"] = email_in.strip().lower()
                    st.success("Logged in as " + st.session_state["user_email"])
                    try:
                        safe_rerun()
                    except Exception:
                        pass
                else:
                    st.error(msg)
    if st.session_state.get("auth_message"):
        st.caption(st.session_state["auth_message"])

if not st.session_state.get("user_email"):
    st.markdown("# Please log in or sign up")
    st.info("Use the sidebar to Login or Sign up with an email and password.")
    st.stop()

# ------------------- Styling -------------------
st.markdown(
    """
    <style>
    .stApp { background-color: #0f1720; color: #e6eef8; }
    .stButton>button, .stFileUploader button, form .stButton>button {
        background-color: #bfbfbf !important;
        color: #000000 !important;
        border-radius: 6px !important;
        border: 1px solid #999 !important;
        font-weight: 600 !important;
        padding: 6px 10px !important;
        font-size: 13px !important;
        height: 34px !important;
    }
    .stButton>button:hover, form .stButton>button:hover, .stFileUploader button:hover {
        background-color: #00b4d8 !important;
        color: #ffffff !important;
    }
    .chat-area { max-height: 480px; overflow-y: auto; padding: 10px; border: 1px solid #333; border-radius: 8px; background-color: #071226; }
    .user-bubble { background: #1e293b; color: #e6eef8; padding: 8px; border-radius: 10px; margin: 6px 40px 6px 6px; white-space: pre-wrap; }
    .bot-bubble { background: #06202e; color: #00ffff; padding: 8px; border-radius: 10px; margin: 6px 6px 6px 40px; white-space: pre-wrap; }
    textarea, input[type="text"] { padding: 8px !important; }
    </style>
    """,
    unsafe_allow_html=True,
)

# ------------------- Main UI -------------------
st.title("🤖 Jarvis — Pinecone + TinyLlama optional")
st.caption("Upload ONE file, ask Jarvis — answers come only from the uploaded file (local or Pinecone).")

st.header("📂 Upload ONE file (PDF / DOCX / TXT)")
uploaded_file = st.file_uploader("Drop a file here (PDF/TXT/DOCX) — one file only", accept_multiple_files=False, type=["pdf", "docx", "txt"])

if uploaded_file is not None:
    # clear previous docs so only this file is indexed
    st.session_state.documents = []
    st.info(f"Processing {uploaded_file.name} ...")
    progress = st.progress(0)
    name = uploaded_file.name
    size = getattr(uploaded_file, "size", None)
    if size is not None:
        st.write(f"**{name}** — {human_readable_size(size)}")
    try:
        data = uploaded_file.read()
        if name.lower().endswith(".pdf"):
            text = extract_text_from_pdf(data)
        elif name.lower().endswith(".docx"):
            text = extract_text_from_docx(data)
        else:
            text = extract_text_from_txt(data)
    except Exception as e:
        st.error(f"Failed to extract from {name}: {e}")
        text = ""
    if not text.strip():
        st.warning(f"No text found in {name}. Skipping.")
    else:
        chunks = chunk_text(text)
        index_chunks_in_session(chunks, source_name=name)
        progress.progress(100)
        st.success("✅ Upload & local indexing complete!")
        # attempt pinecone upsert if enabled
        if PINECONE_ENABLED and SentenceTransformer is not None:
            ok = upsert_to_pinecone(chunks, name)
            if ok:
                st.info("Indexed to Pinecone.")
            else:
                st.info("Pinecone indexing skipped/failed (using local TF-IDF).")

# ------------------- Summarizer -------------------
st.markdown("---")
if st.button("🧾 Summarize Uploaded File (5 lines)"):
    if not st.session_state.documents:
        st.warning("📂 Please upload a file before summarizing.")
    else:
        combined = " ".join(d["text"] for d in st.session_state.documents)
        sentences = re.split(r'(?<=[.!?])\s+', combined)
        seen = set()
        unique_sentences = []
        for s in sentences:
            t = s.strip()
            if not t:
                continue
            if t in seen:
                continue
            seen.add(t)
            unique_sentences.append(t)
        tokens = re.findall(r"\w+", combined.lower())
        freq = {}
        for tok in tokens:
            freq[tok] = freq.get(tok, 0) + 1
        def score_sent(s):
            return sum(freq.get(t.lower(), 0) for t in re.findall(r"\w+", s))
        ranked = sorted(unique_sentences, key=score_sent, reverse=True)
        top5 = [s.strip() for s in ranked if s.strip()][:5]
        if len(top5) < 5:
            for doc in st.session_state.documents:
                txt = doc.get("text","").strip()
                if not txt:
                    continue
                excerpt = txt.splitlines()[0].strip()
                if excerpt and excerpt not in top5:
                    top5.append(excerpt)
                if len(top5) >= 5:
                    break
        final = [ (t if len(t) <= 300 else t[:297].rstrip()+"...") for t in top5[:5] ]
        while len(final) < 5:
            final.append("")
        st.markdown("### 📜 Quick 5-line Summary")
        for line in final:
            st.write("-", line)

# ------------------- Chat input (answers only from uploaded docs) -------------------
st.markdown("---")
st.subheader("💬 Ask Jarvis (answers only from uploaded document)")

with st.form(key="chat_form", clear_on_submit=False):
    user_query = st.text_input("Enter your question here:", key="chat_input")
    cols = st.columns([1, 1, 1])
    with cols[0]:
        send = st.form_submit_button("📤 Send")
    with cols[1]:
        record = st.form_submit_button("🎤 Record (4s)")
    with cols[2]:
        speak_last = st.form_submit_button("🔊 Speak Last")

    if record:
        if not USE_SOUNDDEVICE:
            st.error("Voice recording not available. Install sounddevice & wavio.")
        else:
            try:
                st.info("Recording 4 seconds... Speak clearly")
                tmp_wav = record_with_sounddevice(seconds=4)
                import speech_recognition as sr
                r = sr.Recognizer()
                with sr.AudioFile(str(tmp_wav)) as source:
                    audio = r.record(source)
                try:
                    query_text = r.recognize_google(audio)
                    st.session_state['_last_recorded'] = query_text
                    st.success(f"You said: {query_text}")
                    st.session_state["chat_input"] = query_text
                except sr.UnknownValueError:
                    st.error("Could not understand audio")
                except sr.RequestError:
                    st.error("Speech recognition failed (network)")
                finally:
                    tmp_wav.unlink(missing_ok=True)
            except Exception as e:
                st.error(f"Recording failed: {e}")

    if speak_last:
        if st.session_state.chat_history:
            last = st.session_state.chat_history[-1]["bot"]
            try:
                if USE_PYTTSX3:
                    threaded_tts_play(last)
                else:
                    st.info("TTS not available.")
            except Exception:
                pass
        else:
            st.info("No answer yet to speak.")

    if send and user_query and user_query.strip():
        q = user_query.strip()
        if not st.session_state.documents:
            st.warning("📂 Please upload a file before asking questions.")
        else:
            # robust submission token to avoid duplicate processing on rerun
            if not st.session_state.get("_cur_submission_token"):
                st.session_state["_cur_submission_token"] = uuid.uuid4().hex
            cur_token = st.session_state["_cur_submission_token"]
            if st.session_state.get("_handled_submission_token") == cur_token:
                pass
            else:
                # prefer Pinecone if enabled and embedder available
                answers = []
                if PINECONE_ENABLED and SentenceTransformer is not None:
                    try:
                        pc_results = query_pinecone(q, top_k=6)
                        for r in pc_results:
                            answers.append({"source": r.get("source", "unknown"), "text": r.get("text", "")})
                    except Exception:
                        answers = []
                # if no pinecone results or pinecone not enabled, fallback to TF-IDF
                if not answers:
                    tf_docs = retrieve_from_session_tfidf(q, k=6)
                    answers = [{"source": d.get("source","unknown"), "text": d.get("text","")} for d in tf_docs]

                if not answers:
                    answer_text = "I don't know — I couldn't find relevant information in the uploaded document for that question. Please check the uploaded file or ask a different question."
                else:
                    # dedupe by source
                    seen_src = set()
                    deduped = []
                    for a in answers:
                        if a["source"] in seen_src:
                            continue
                        seen_src.add(a["source"])
                        deduped.append(a)
                        if len(deduped) >= 3:
                            break

                    # sentence scoring
                    query_lower = q.lower()
                    query_tokens = set(re.findall(r"\w+", query_lower))
                    scored_sentences = []
                    for d in deduped:
                        txt = d["text"].replace("\n", " ")
                        sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', txt) if s.strip()]
                        best = (0.0, "")
                        for s in sentences:
                            s_lower = s.lower()
                            tokens = re.findall(r"\w+", s_lower)
                            if not tokens:
                                continue
                            shared = sum(1 for t in set(tokens) if t in query_tokens)
                            score = shared / (math.sqrt(len(tokens)) + 1e-6)
                            if query_lower in s_lower:
                                score += 0.5
                            if score > best[0]:
                                best = (score, s.strip())
                        if best[0] > 0:
                            scored_sentences.append((best[0], best[1], d["source"]))
                    scored_sentences.sort(key=lambda x: x[0], reverse=True)
                    if not scored_sentences:
                        answer_text = "I don't know — I couldn't find relevant information in the uploaded document for that question. Please check the uploaded file or ask a different question."
                    else:
                        # detect list style queries
                        triggers = ["type", "types", "what are", "examples", "list", "kinds of", "give me", "define", "difference"]
                        is_list = any(t in query_lower for t in triggers)
                        if is_list:
                            seen_snips = set()
                            bullets = []
                            for _, s, src in scored_sentences:
                                snippet = s if len(s) <= 280 else s[:277].rstrip() + "..."
                                norm = re.sub(r'\s+', ' ', snippet).strip().lower()
                                if norm in seen_snips:
                                    continue
                                seen_snips.add(norm)
                                bullets.append(f"- {snippet}")
                                if len(bullets) >= 5:
                                    break
                            sources = ", ".join(sorted({d["source"] for d in deduped}))[:200]
                            answer_text = "Here are the main items found in the document:\n\n" + "\n".join(bullets) + f"\n\n(sources: {sources})"
                        else:
                            top_sentences = []
                            used = set()
                            for _, s, src in scored_sentences:
                                norm = s.strip()
                                if norm in used:
                                    continue
                                top_sentences.append(norm)
                                used.add(norm)
                                if len(top_sentences) >= 2:
                                    break
                            concise = make_concise_paragraph(top_sentences, max_sentences=3, max_chars=420)
                            sources = ", ".join(sorted({d["source"] for d in deduped}))[:200]
                            answer_text = f"{concise} (sources: {sources})"

                # append to history guarded by token
                if st.session_state.get("last_processed_query") != q or st.session_state.get("_handled_submission_token") != cur_token:
                    st.session_state.chat_history.append({"user": q, "bot": answer_text, "timestamp": int(time.time())})
                    st.session_state.last_processed_query = q
                    st.session_state["_handled_submission_token"] = cur_token
                    st.session_state["_cur_submission_token"] = None
                    try:
                        if USE_PYTTSX3:
                            threaded_tts_play(answer_text)
                    except Exception:
                        pass
                    try:
                        st.session_state["chat_input"] = ""
                    except Exception:
                        pass
                safe_rerun()

# ------------------- Chat area -------------------
st.markdown('<div class="chat-area">', unsafe_allow_html=True)
for c in st.session_state.chat_history:
    st.markdown(f"<div class='user-bubble'>🧑‍💻 {c['user']}</div>", unsafe_allow_html=True)
    st.markdown(f"<div class='bot-bubble'>🤖 {c['bot']}</div>", unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)

# ------------------- Export chat to PDF / DOCX -------------------
st.markdown("---")
col_a, col_b = st.columns(2)
with col_a:
    if st.button("📤 Export Chat to PDF"):
        if not st.session_state.chat_history:
            st.info("No chat to export")
        else:
            pdf_path = Path(tempfile.gettempdir()) / f"Jarvis_Chat_{uuid.uuid4().hex}.pdf"
            try:
                if FPDF is None:
                    st.error("FPDF not installed (pip install fpdf).")
                else:
                    pdf = FPDF()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.add_page()
                    pdf.set_font("Arial", size=12)
                    pdf.cell(0, 8, safe_for_pdf("Jarvis Chat History"), ln=True)
                    pdf.ln(4)
                    for c in st.session_state.chat_history:
                        ts = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(c.get('timestamp', time.time())))
                        pdf.multi_cell(0, 7, safe_for_pdf(f"[{ts}] You: {c.get('user','')}"))
                        pdf.multi_cell(0, 7, safe_for_pdf(f"Jarvis: {c.get('bot','')}"))
                        pdf.ln(2)
                    pdf.output(str(pdf_path))
                    with open(pdf_path, "rb") as fh:
                        st.download_button("⬇️ Download PDF", fh.read(), file_name="Jarvis_Chat.pdf")
            except Exception as e:
                st.error("PDF export failed: " + str(e))

with col_b:
    if st.button("📄 Export Chat to Word"):
        if not st.session_state.chat_history:
            st.info("No chat to export")
        else:
            doc_path = Path(tempfile.gettempdir()) / f"Jarvis_Chat_{uuid.uuid4().hex}.docx"
            try:
                if DocxDocument is None:
                    st.error("python-docx not installed (pip install python-docx).")
                else:
                    doc = DocxDocument()
                    doc.add_heading("Jarvis Chat History", level=1)
                    for c in st.session_state.chat_history:
                        ts = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(c.get('timestamp', time.time())))
                        doc.add_paragraph(f"[{ts}] You: {c.get('user','')}")
                        doc.add_paragraph(f"Jarvis: {c.get('bot','')}")
                        doc.add_paragraph("")
                    doc.save(str(doc_path))
                    with open(doc_path, "rb") as fh:
                        st.download_button("⬇️ Download DOCX", fh.read(), file_name="Jarvis_Chat.docx")
            except Exception as e:
                st.error("DOCX export failed: " + str(e))

# Footer status
st.markdown("---")
status_parts = []
if PINECONE_ENABLED:
    status_parts.append("Pinecone: ✅")
else:
    status_parts.append("Pinecone: disabled (fallback TF-IDF)")

if LLM_ENABLED:
    status_parts.append(f"LLM: {OLLAMA_MODEL} ✅")
else:
    status_parts.append("LLM: disabled (no Ollama client)")

st.caption(" | ".join(status_parts) + f" — Signed in as: {st.session_state.get('user_email')}")
# app_with_uploader_pro_with_auth.py
"""
Jarvis — Streamlit document assistant (local, concise) with simple email/password auth.
Features: signup/login (local sqlite) -> upload (PDF/TXT/DOCX) -> local in-memory indexing,
chat (answers ONLY from uploaded documents), summarizer, voice input (optional), TTS (optional),
export to PDF/DOCX, New Chat (clears history).

Security note: this auth is suitable for local/dev only (stores PBKDF2-hashed passwords in users.db).
"""

import os
from pathlib import Path
import uuid
import io
import threading
import tempfile
import time
import re
import unicodedata
import math
import sqlite3
import hashlib
import binascii

import streamlit as st

# ------------------- dotenv (optional) ----------------
from dotenv import load_dotenv
env_path = Path(__file__).resolve().parent / ".env"
try:
    load_dotenv(dotenv_path=env_path)
except Exception:
    pass

# ------------------- lightweight imports -------------------
import PyPDF2
import docx

# optional voice libraries - attempt imports and toggle features
USE_SOUNDDEVICE = False
USE_PYTTSX3 = False
try:
    import sounddevice as sd
    import wavio
    USE_SOUNDDEVICE = True
except Exception:
    USE_SOUNDDEVICE = False

try:
    import pyttsx3
    USE_PYTTSX3 = True
except Exception:
    USE_PYTTSX3 = False

# ------------------- Streamlit page configuration -------------------
# Must be called before any other Streamlit UI calls
st.set_page_config(page_title="Jarvis (Local Docs Only)", layout="wide", initial_sidebar_state="expanded")

# ------------------- helpers (needed for auth UI) -------------------
def safe_rerun():
    # Compatible rerun helper
    if hasattr(st, "rerun"):
        try:
            st.rerun()
            return
        except Exception:
            pass
    if hasattr(st, "experimental_rerun"):
        try:
            st.experimental_rerun()
            return
        except Exception:
            pass
    st.session_state["_force_update_flag"] = not st.session_state.get("_force_update_flag", False)

def human_readable_size(num_bytes: int) -> str:
    for unit in ["B", "KB", "MB", "GB", "TB"]:
        if abs(num_bytes) < 1024.0:
            return f"{num_bytes:3.1f} {unit}"
        num_bytes /= 1024.0
    return f"{num_bytes:.1f} PB"

def safe_for_pdf(s: str) -> str:
    """Return a latin1-safe version of `s` for fpdf (drops non-latin1 chars)."""
    if s is None:
        return ""
    if not isinstance(s, str):
        s = str(s)
    normalized = unicodedata.normalize("NFKD", s)
    return normalized.encode("latin-1", "ignore").decode("latin-1")

# ------------------- Simple email/password auth (sqlite + PBKDF2) -------------------
USERS_DB_PATH = Path(__file__).resolve().parent / "users.db"

def init_user_db():
    conn = sqlite3.connect(str(USERS_DB_PATH))
    cur = conn.cursor()
    cur.execute("""
      CREATE TABLE IF NOT EXISTS users (
        email TEXT PRIMARY KEY,
        salt TEXT NOT NULL,
        pw_hash TEXT NOT NULL,
        created_at INTEGER NOT NULL
      )
    """)
    conn.commit()
    conn.close()

def hash_password(password: str, salt: bytes = None):
    """
    Return (salt_hex, hash_hex). Uses PBKDF2-HMAC-SHA256.
    """
    if salt is None:
        salt = os.urandom(16)
    dk = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, 100_000)
    return binascii.hexlify(salt).decode("ascii"), binascii.hexlify(dk).decode("ascii")

def create_user(email: str, password: str) -> (bool, str):
    email = (email or "").strip().lower()
    if not email or not password:
        return False, "Email and password required."
    conn = sqlite3.connect(str(USERS_DB_PATH))
    cur = conn.cursor()
    cur.execute("SELECT 1 FROM users WHERE email = ?", (email,))
    if cur.fetchone():
        conn.close()
        return False, "Email already registered."
    salt_hex, hash_hex = hash_password(password)
    ts = int(time.time())
    try:
        cur.execute("INSERT INTO users (email, salt, pw_hash, created_at) VALUES (?, ?, ?, ?)",
                    (email, salt_hex, hash_hex, ts))
        conn.commit()
    except Exception as e:
        conn.close()
        return False, f"DB error: {e}"
    conn.close()
    return True, "Account created."

def verify_user(email: str, password: str) -> (bool, str):
    email = (email or "").strip().lower()
    conn = sqlite3.connect(str(USERS_DB_PATH))
    cur = conn.cursor()
    cur.execute("SELECT salt, pw_hash FROM users WHERE email = ?", (email,))
    row = cur.fetchone()
    conn.close()
    if not row:
        return False, "No account with that email."
    salt_hex, stored_hash_hex = row
    salt = binascii.unhexlify(salt_hex.encode("ascii"))
    _, hash_hex = hash_password(password, salt=salt)
    if hash_hex == stored_hash_hex:
        return True, "Login successful."
    return False, "Incorrect password."

# Initialize DB and session keys
init_user_db()
if "user_email" not in st.session_state:
    st.session_state["user_email"] = None
if "auth_message" not in st.session_state:
    st.session_state["auth_message"] = ""

# ------------------- If not logged in show auth UI and exit early -------------------
with st.sidebar:
    st.sidebar.title("🔐 Account")
    if st.session_state.get("user_email"):
        st.markdown(f"**Signed in as:** {st.session_state['user_email']}")
        if st.button("🔓 Logout"):
            st.session_state["user_email"] = None
            st.session_state["auth_message"] = "Logged out."
            try:
                safe_rerun()
            except Exception:
                pass
    else:
        auth_mode = st.radio("Choose:", ["Login", "Sign up"], index=0)
        email_in = st.text_input("Email", key="auth_email")
        pw_in = st.text_input("Password", type="password", key="auth_pw")
        if auth_mode == "Sign up":
            if st.button("Create account"):
                ok, msg = create_user(email_in, pw_in)
                st.session_state["auth_message"] = msg
                if ok:
                    st.success(msg)
                else:
                    st.error(msg)
        else:  # Login
            if st.button("Login"):
                ok, msg = verify_user(email_in, pw_in)
                st.session_state["auth_message"] = msg
                if ok:
                    st.session_state["user_email"] = email_in.strip().lower()
                    st.success("Logged in as " + st.session_state["user_email"])
                    try:
                        safe_rerun()
                    except Exception:
                        pass
                else:
                    st.error(msg)

    if st.session_state.get("auth_message"):
        st.caption(st.session_state["auth_message"])

# If not authenticated, do not render the main app UI
if not st.session_state.get("user_email"):
    st.markdown("# Please log in or sign up")
    st.info("Use the sidebar to Login or Sign up with an email and password.")
    st.stop()

# ------------------- Styling: concise buttons, compact UI -------------------
st.markdown(
    """
    <style>
    .stApp { background-color: #0f1720; color: #e6eef8; }

    /* concise buttons */
    .stButton>button, .stFileUploader button, form .stButton>button {
        background-color: #bfbfbf !important;   /* gray background */
        color: #000000 !important;               /* black text */
        border-radius: 6px !important;
        border: 1px solid #999 !important;
        font-weight: 600 !important;
        padding: 6px 10px !important;
        font-size: 13px !important;
        height: 34px !important;
    }
    .stButton>button:hover, form .stButton>button:hover, .stFileUploader button:hover {
        background-color: #00b4d8 !important;    /* blue hover */
        color: #ffffff !important;
    }

    /* compact file uploader text */
    .stFileUploader { padding: 6px 0 !important; }

    /* sidebar */
    section[data-testid="stSidebar"] { background-color: #000000; color: #e6eef8; border-right: 1px solid #333; }
    section[data-testid="stSidebar"] h1, section[data-testid="stSidebar"] h2 { color: #00ffff !important; }

    /* chat area */
    .chat-area { max-height: 480px; overflow-y: auto; padding: 10px; border: 1px solid #333; border-radius: 8px; background-color: #071226; }
    .user-bubble { background: #1e293b; color: #e6eef8; padding: 8px; border-radius: 10px; margin: 6px 40px 6px 6px; white-space: pre-wrap; }
    .bot-bubble { background: #06202e; color: #00ffff; padding: 8px; border-radius: 10px; margin: 6px 6px 6px 40px; white-space: pre-wrap; }

    /* make text input compact */
    textarea, input[type="text"] { padding: 8px !important; }
    </style>
    """,
    unsafe_allow_html=True,
)

# ------------------- REPLACED: improved make_concise_paragraph -------------------
def make_concise_paragraph(snippets, max_sentences=3, max_chars=420):
    if not snippets:
        return ""
    merged = []
    for s in snippets:
        if not s:
            continue
        t = s.strip().replace("\n", " ")
        t = re.sub(r"[“”«»]", '"', t)
        t = re.sub(r"[‘’]", "'", t)
        if len(t) < 3:
            continue
        merged.append(t)
    if not merged:
        return ""
    first = merged[0]
    extend_needed = False
    if first.endswith(":") or first.endswith("?") or len(first.split()) <= 6:
        extend_needed = True
    if extend_needed and len(merged) > 1:
        combined = " ".join(merged[: min(len(merged), 4)])
    else:
        combined = " ".join(merged)
    sentences = re.split(r'(?<=[.!?])\s+', combined)
    sentences = [s.strip() for s in sentences if s.strip()]
    if not sentences:
        paragraph = combined.strip()
    else:
        paragraph = " ".join(sentences[:max_sentences])
    paragraph = re.sub(r"•\s*", "", paragraph)
    paragraph = re.sub(r"\s+", " ", paragraph).strip()
    if len(paragraph) > max_chars:
        paragraph = paragraph[: max_chars - 3].rstrip() + "..."
    return paragraph

# ------------------- session state init -------------------
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "documents" not in st.session_state:
    st.session_state.documents = []
if "last_processed_query" not in st.session_state:
    st.session_state.last_processed_query = None
if "_cur_submission_token" not in st.session_state:
    st.session_state["_cur_submission_token"] = None
if "_handled_submission_token" not in st.session_state:
    st.session_state["_handled_submission_token"] = None

# ------------------- parsers -------------------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    # Improved merge of lines to avoid headings-only chunks
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    pages = []
    for p in reader.pages:
        try:
            raw = p.extract_text() or ""
        except Exception:
            raw = ""
        lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
        merged = []
        buffer = ""
        for line in lines:
            if re.search(r'[.!?]$', line):
                buffer += " " + line
                merged.append(buffer.strip())
                buffer = ""
            else:
                buffer += " " + line
        if buffer.strip():
            merged.append(buffer.strip())
        pages.append(" ".join(merged))
    text = "\n".join(pages)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'•', '-', text)
    return text.strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    tmp = Path(tempfile.gettempdir()) / f"tmp_{uuid.uuid4().hex}.docx"
    tmp.write_bytes(file_bytes)
    doc = docx.Document(str(tmp))
    paragraphs = [p.text for p in doc.paragraphs]
    tmp.unlink(missing_ok=True)
    return "\n".join(paragraphs)

def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")

def chunk_text(text: str, chunk_size: int = 800, overlap: int = 150):
    chunks = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + chunk_size, length)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start += chunk_size - overlap
    return chunks

def index_chunks_in_session(chunks, source_name):
    for c in chunks:
        st.session_state.documents.append({"source": source_name, "text": c})

# ---------- TF-IDF index builder + retriever ----------
def build_tfidf_index():
    docs = st.session_state.documents
    N = len(docs)
    if N == 0:
        return [], {}
    doc_tokens = []
    df = {}
    for d in docs:
        tokens = re.findall(r"\w+", d["text"].lower())
        uniq = set(tokens)
        for t in uniq:
            df[t] = df.get(t, 0) + 1
        doc_tokens.append(tokens)
    idf = {}
    for t, cnt in df.items():
        idf[t] = math.log((N + 1) / (cnt + 1)) + 1.0
    docs_vectors = []
    for tokens, d in zip(doc_tokens, docs):
        tf = {}
        for t in tokens:
            tf[t] = tf.get(t, 0) + 1
        vec = {}
        norm_sq = 0.0
        for t, tfcount in tf.items():
            w = tfcount * idf.get(t, 0.0)
            vec[t] = w
            norm_sq += w * w
        norm = math.sqrt(norm_sq) if norm_sq > 0 else 1.0
        docs_vectors.append({"doc": d, "vec": vec, "norm": norm})
    return docs_vectors, idf

def retrieve_from_session(query: str, k: int = 3):
    if not st.session_state.documents or not query or not query.strip():
        return []
    q_tokens = re.findall(r"\w+", query.lower())
    if not q_tokens:
        return []
    docs_vectors, idf = build_tfidf_index()
    if not docs_vectors:
        return []
    q_tf = {}
    for t in q_tokens:
        q_tf[t] = q_tf.get(t, 0) + 1
    q_vec = {}
    q_norm_sq = 0.0
    for t, tfcount in q_tf.items():
        idf_val = idf.get(t, math.log((len(docs_vectors) + 1) / 1) + 1.0)
        w = tfcount * idf_val
        q_vec[t] = w
        q_norm_sq += w * w
    q_norm = math.sqrt(q_norm_sq) if q_norm_sq > 0 else 1.0
    scored = []
    q_text_lower = query.lower()
    for item in docs_vectors:
        doc = item["doc"]
        doc_vec = item["vec"]
        dot = 0.0
        for t, qw in q_vec.items():
            dw = doc_vec.get(t)
            if dw:
                dot += qw * dw
        denom = (q_norm * item["norm"]) if (q_norm and item["norm"]) else 1.0
        score = dot / denom if denom != 0 else 0.0
        if q_text_lower in doc["text"].lower():
            score += 0.18
        common_tokens = sum(1 for t in set(q_tokens) if t in doc["text"].lower())
        score += min(common_tokens * 0.02, 0.06)
        scored.append((score, doc))
    scored.sort(key=lambda x: x[0], reverse=True)
    top = []
    for s, d in scored:
        if s <= 0:
            continue
        top.append(d)
        if len(top) >= k:
            break
    return top

# Threaded TTS helper
def threaded_tts_play(text: str):
    if USE_PYTTSX3:
        def _speak(t):
            try:
                engine = pyttsx3.init()
                engine.say(t)
                engine.runAndWait()
            except Exception:
                pass
        threading.Thread(target=_speak, args=(text,), daemon=True).start()

# voice recording helper
def record_with_sounddevice(seconds=4, fs=44100):
    if not USE_SOUNDDEVICE:
        raise RuntimeError("sounddevice not available")
    recording = sd.rec(int(seconds * fs), samplerate=fs, channels=1)
    sd.wait()
    tmp_wav = Path(tempfile.gettempdir()) / f"rec_{uuid.uuid4().hex}.wav"
    wavio.write(str(tmp_wav), recording, fs, sampwidth=2)
    return tmp_wav

# ------------------- Sidebar: New Chat & history -------------------
st.sidebar.title("💬 Controls")
if st.sidebar.button("➕ New Chat"):
    st.session_state.chat_history = []
    st.session_state.documents = []
    st.session_state.last_processed_query = None
    st.session_state["_cur_submission_token"] = None
    st.session_state["_handled_submission_token"] = None
    safe_rerun()

st.sidebar.markdown("### Recent conversation")
for chat in reversed(st.session_state.chat_history[-30:]):
    st.sidebar.markdown(f"**You:** {chat['user']}")
    st.sidebar.markdown(f"**Jarvis:** {chat['bot']}")
    st.sidebar.markdown("---")

# ------------------- Main UI: upload (single file) -------------------
st.title("🤖 Jarvis — Local Document Assistant")
st.caption("Upload documents and ask Jarvis. Jarvis will answer ONLY from uploaded documents.")

st.header("📂 Upload ONE file (PDF / DOCX / TXT)")
uploaded_file = st.file_uploader("Drop a file here (PDF/TXT/DOCX) — one file only", accept_multiple_files=False, type=["pdf", "docx", "txt"])

if uploaded_file is not None:
    st.session_state.documents = []
    st.info(f"Processing {uploaded_file.name} ...")
    progress = st.progress(0)
    name = uploaded_file.name
    size = getattr(uploaded_file, "size", None)
    if size is not None:
        st.write(f"**{name}** — {human_readable_size(size)}")
    try:
        data = uploaded_file.read()
        if name.lower().endswith(".pdf"):
            text = extract_text_from_pdf(data)
        elif name.lower().endswith(".docx"):
            text = extract_text_from_docx(data)
        else:
            text = extract_text_from_txt(data)
    except Exception as e:
        st.error(f"Failed to extract from {name}: {e}")
        text = ""
    if not text.strip():
        st.warning(f"No text found in {name}. Skipping.")
    else:
        chunks = chunk_text(text)
        index_chunks_in_session(chunks, source_name=name)
        progress.progress(100)
        st.success("✅ Upload & local indexing complete!")

# ------------------- Summarizer (local) — produce exactly 5 clear lines -------------------
st.markdown("---")
if st.button("🧾 Summarize Uploaded Files (local)"):
    if not st.session_state.documents:
        st.warning("📂 Please upload at least one document before summarizing.")
    else:
        combined = " ".join(d["text"] for d in st.session_state.documents)
        sentences = re.split(r'(?<=[.!?])\s+', combined)
        seen = set()
        unique_sentences = []
        for s in sentences:
            t = s.strip()
            if not t:
                continue
            if t in seen:
                continue
            seen.add(t)
            unique_sentences.append(t)
        tokens = re.findall(r"\w+", combined.lower())
        freq = {}
        for tok in tokens:
            freq[tok] = freq.get(tok, 0) + 1
        def score_sent(s):
            return sum(freq.get(t.lower(), 0) for t in re.findall(r"\w+", s))
        ranked = sorted(unique_sentences, key=score_sent, reverse=True)
        top5 = [s.strip() for s in ranked if s.strip()][:5]
        if len(top5) < 5:
            for doc in st.session_state.documents:
                txt = doc.get("text","").strip()
                if not txt:
                    continue
                excerpt = txt.splitlines()[0].strip()
                if excerpt and excerpt not in top5:
                    top5.append(excerpt)
                if len(top5) >= 5:
                    break
        final = [ (t if len(t) <= 300 else t[:297].rstrip()+"...") for t in top5[:5] ]
        while len(final) < 5:
            final.append("")
        st.markdown("### 📜 Quick 5-line Summary")
        for line in final:
            st.write("-", line)

# ------------------- Chat input (answers only from uploaded docs) -------------------
st.markdown("---")
st.subheader("💬 Ask Jarvis (answers only from uploaded documents)")

with st.form(key="chat_form", clear_on_submit=False):
    user_query = st.text_input("Enter your question here:", key="chat_input")
    cols = st.columns([1, 1, 1])
    with cols[0]:
        send = st.form_submit_button("📤 Send")
    with cols[1]:
        record = st.form_submit_button("🎤 Record (4s)")
    with cols[2]:
        speak_last = st.form_submit_button("🔊 Speak Last")

    if record:
        if not USE_SOUNDDEVICE:
            st.error("Voice recording not available. Install sounddevice & wavio.")
        else:
            try:
                st.info("Recording 4 seconds... Speak clearly")
                tmp_wav = record_with_sounddevice(seconds=4)
                import speech_recognition as sr
                r = sr.Recognizer()
                with sr.AudioFile(str(tmp_wav)) as source:
                    audio = r.record(source)
                try:
                    query_text = r.recognize_google(audio)
                    st.session_state['_last_recorded'] = query_text
                    st.success(f"You said: {query_text}")
                    st.session_state["chat_input"] = query_text
                except sr.UnknownValueError:
                    st.error("Could not understand audio")
                except sr.RequestError:
                    st.error("Speech recognition failed (network)")
                finally:
                    tmp_wav.unlink(missing_ok=True)
            except Exception as e:
                st.error(f"Recording failed: {e}")

    if speak_last:
        if st.session_state.chat_history:
            last = st.session_state.chat_history[-1]["bot"]
            try:
                threaded_tts_play(last)
            except Exception:
                pass
        else:
            st.info("No answer yet to speak.")

    if send and user_query and user_query.strip():
        q = user_query.strip()
        if not st.session_state.documents:
            st.warning("📂 Please upload at least one document before asking questions. Please upload file first.")
        else:
            if not st.session_state.get("_cur_submission_token"):
                st.session_state["_cur_submission_token"] = uuid.uuid4().hex
            cur_token = st.session_state["_cur_submission_token"]
            if st.session_state.get("_handled_submission_token") == cur_token:
                pass
            else:
                top_docs = retrieve_from_session(q, k=6)
                if not top_docs:
                    answer = "I don't know — I couldn't find relevant information in the uploaded document for that question. Please check the uploaded file or ask a different question."
                else:
                    seen_src = set()
                    deduped_docs = []
                    for d in top_docs:
                        if d["source"] in seen_src:
                            continue
                        seen_src.add(d["source"])
                        deduped_docs.append(d)
                        if len(deduped_docs) >= 3:
                            break
                    def is_list_request(query_lower: str) -> bool:
                        triggers = ["type", "types", "what are", "examples", "list", "kinds of", "give me", "define", "difference"]
                        return any(t in query_lower for t in triggers)
                    query_lower = q.lower()
                    query_tokens = set(re.findall(r"\w+", query_lower))
                    scored_sentences = []
                    for d in deduped_docs:
                        txt = d["text"].replace("\n", " ")
                        sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', txt) if s.strip()]
                        best = (0.0, "")
                        for s in sentences:
                            s_lower = s.lower()
                            tokens = re.findall(r"\w+", s_lower)
                            if not tokens:
                                continue
                            shared = sum(1 for t in set(tokens) if t in query_tokens)
                            score = shared / (math.sqrt(len(tokens)) + 1e-6)
                            if query_lower in s_lower:
                                score += 0.5
                            if score > best[0]:
                                best = (score, s.strip())
                        if best[0] > 0:
                            scored_sentences.append((best[0], best[1], d["source"]))
                    scored_sentences.sort(key=lambda x: x[0], reverse=True)
                    if not scored_sentences:
                        answer = "I don't know — I couldn't find relevant information in the uploaded document for that question. Please check the uploaded file or ask a different question."
                    else:
                        if is_list_request(query_lower):
                            seen_snips = set()
                            bullets = []
                            for _, s, src in scored_sentences:
                                snippet = s if len(s) <= 280 else s[:277].rstrip() + "..."
                                norm = re.sub(r'\s+', ' ', snippet).strip().lower()
                                if norm in seen_snips:
                                    continue
                                seen_snips.add(norm)
                                bullets.append(f"- {snippet}")
                                if len(bullets) >= 5:
                                    break
                            sources = ", ".join(sorted({d["source"] for d in deduped_docs}))[:200]
                            answer = "Here are the main items found in the document:\n\n" + "\n".join(bullets) + f"\n\n(sources: {sources})"
                        else:
                            top_sentences = []
                            used = set()
                            for _, s, src in scored_sentences:
                                norm = s.strip()
                                if norm in used:
                                    continue
                                top_sentences.append(norm)
                                used.add(norm)
                                if len(top_sentences) >= 2:
                                    break
                            concise = make_concise_paragraph(top_sentences, max_sentences=3, max_chars=420)
                            sources = ", ".join(sorted({d["source"] for d in deduped_docs}))[:200]
                            answer = f"{concise} (sources: {sources})"
                if st.session_state.get("last_processed_query") != q or st.session_state.get("_handled_submission_token") != cur_token:
                    st.session_state.chat_history.append({"user": q, "bot": answer, "timestamp": int(time.time())})
                    st.session_state.last_processed_query = q
                    st.session_state["_handled_submission_token"] = cur_token
                    st.session_state["_cur_submission_token"] = None
                    try:
                        threaded_tts_play(answer)
                    except Exception:
                        pass
                    try:
                        st.session_state["chat_input"] = ""
                    except Exception:
                        pass
                safe_rerun()

# ------------------- Chat history area (single display) ---
st.markdown('<div class="chat-area">', unsafe_allow_html=True)
for c in st.session_state.chat_history:
    st.markdown(f"<div class='user-bubble'>🧑‍💻 {c['user']}</div>", unsafe_allow_html=True)
    st.markdown(f"<div class='bot-bubble'>🤖 {c['bot']}</div>", unsafe_allow_html=True)
st.markdown('</div>', unsafe_allow_html=True)

# ------------------- Export chat to PDF or Word -------------------
st.markdown("---")
col_a, col_b = st.columns(2)
with col_a:
    if st.button("📤 Export Chat to PDF"):
        if not st.session_state.chat_history:
            st.info("No chat to export")
        else:
            pdf_path = Path(tempfile.gettempdir()) / f"Jarvis_Chat_{uuid.uuid4().hex}.pdf"
            try:
                from fpdf import FPDF
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.cell(0, 8, safe_for_pdf("Jarvis Chat History"), ln=True)
                pdf.ln(4)
                for c in st.session_state.chat_history:
                    ts = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(c.get('timestamp', time.time())))
                    pdf.multi_cell(0, 7, safe_for_pdf(f"[{ts}] You: {c.get('user','')}"))
                    pdf.multi_cell(0, 7, safe_for_pdf(f"Jarvis: {c.get('bot','')}"))
                    pdf.ln(2)
                pdf.output(str(pdf_path))
                with open(pdf_path, "rb") as fh:
                    st.download_button("⬇️ Download PDF", fh.read(), file_name="Jarvis_Chat.pdf")
            except Exception as e:
                st.error("PDF export failed: " + str(e))

with col_b:
    if st.button("📄 Export Chat to Word"):
        if not st.session_state.chat_history:
            st.info("No chat to export")
        else:
            doc_path = Path(tempfile.gettempdir()) / f"Jarvis_Chat_{uuid.uuid4().hex}.docx"
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument()
                doc.add_heading("Jarvis Chat History", level=1)
                for c in st.session_state.chat_history:
                    ts = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(c.get('timestamp', time.time())))
                    doc.add_paragraph(f"[{ts}] You: {c.get('user','')}")
                    doc.add_paragraph(f"Jarvis: {c.get('bot','')}")
                    doc.add_paragraph("")
                doc.save(str(doc_path))
                with open(doc_path, "rb") as fh:
                    st.download_button("⬇️ Download DOCX", fh.read(), file_name="Jarvis_Chat.docx")
            except Exception as e:
                st.error("DOCX export failed: " + str(e))

# ------------------- Footer -------------------
st.markdown("---")
st.caption(f"Signed in as: {st.session_state.get('user_email')} — Tip: This lightweight version answers only from uploaded documents (no Pinecone / external model).")

